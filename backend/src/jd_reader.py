"""
JD Reader
Reads job description from .docx format and returns clean text.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def read_jd_docx(file_path: str) -> str:
    """
    Read a .docx job description file and return clean text.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"JD file not found: {file_path}")

    if not path.suffix.lower() == ".docx":
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    doc = Document(str(path))

    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract tables (some JDs use tables for requirements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in text_parts:
                    text_parts.append(text)

    full_text = "\n".join(text_parts)

    logger.info(
        f"  Read JD from {path.name}: {len(full_text)} characters, "
        f"{len(text_parts)} paragraphs"
    )

    if len(full_text) < 50:
        logger.warning("  JD text seems very short — check the .docx file!")

    return full_text


def read_redrob_signals_doc(file_path: str) -> str:
    """Read the redrob signals documentation."""
    try:
        from docx import Document

        path = Path(file_path)
        if not path.exists():
            logger.warning(f"Redrob signals doc not found: {file_path}")
            return ""

        doc = Document(str(path))
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    except Exception as e:
        logger.warning(f"Could not read redrob signals doc: {e}")
        return ""